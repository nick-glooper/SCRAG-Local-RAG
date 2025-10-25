"""
Document processing pipeline for SecureRAG.
Handles extraction, chunking, embedding, and storage.
"""

import logging
import time
from pathlib import Path
from typing import Dict, Optional, Tuple
from datetime import datetime
import hashlib

# Document parsers
import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Orchestrate the complete document ingestion pipeline.
    """

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.doc'}

    def __init__(self, config, vector_store, embedding_handler, chunker):
        """
        Initialize document processor with dependencies.

        Args:
            config: SecureRAGConfig instance
            vector_store: VectorStore instance
            embedding_handler: EmbeddingHandler instance
            chunker: SemanticChunker instance
        """
        self.config = config
        self.vector_store = vector_store
        self.embedding_handler = embedding_handler
        self.chunker = chunker

        logger.info("Document processor initialized")

    def ingest_document(
        self,
        filepath: str,
        collection_name: str,
        metadata: Optional[Dict] = None,
        document_id: Optional[str] = None,
        version: str = "v1"
    ) -> Dict:
        """
        Complete ingestion pipeline for a document.

        Args:
            filepath: Path to document file
            collection_name: Target collection
            metadata: Optional metadata to attach
            document_id: Optional document ID (auto-generated if not provided)
            version: Document version string

        Returns:
            Dict with ingestion results and statistics
        """
        start_time = time.time()

        try:
            # 1. Validation
            file_path = Path(filepath)
            if not file_path.exists():
                return {
                    "success": False,
                    "error": f"File not found: {filepath}",
                    "error_type": "FileNotFoundError"
                }

            if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_path.suffix}",
                    "error_type": "UnsupportedFileTypeError",
                    "supported_types": list(self.SUPPORTED_EXTENSIONS)
                }

            # Generate document ID if not provided
            if document_id is None:
                document_id = self._generate_document_id(file_path)

            logger.info(f"Ingesting document: {file_path.name} -> {collection_name}")

            # 2. Extract text and metadata
            text, extracted_metadata = self._extract_document(file_path)

            if not text or not text.strip():
                return {
                    "success": False,
                    "error": "No text could be extracted from document",
                    "error_type": "EmptyDocumentError"
                }

            # Merge metadata
            full_metadata = {
                **extracted_metadata,
                **(metadata or {}),
                "filename": file_path.name,
                "filepath": str(file_path),
                "file_size_bytes": file_path.stat().st_size,
                "ingested_at": datetime.utcnow().isoformat()
            }

            # 3. Chunk document
            logger.info(f"Chunking document ({len(text)} characters)")
            chunks = self.chunker.chunk_document(
                text,
                full_metadata,
                filename=file_path.name
            )

            if not chunks:
                return {
                    "success": False,
                    "error": "No chunks generated from document",
                    "error_type": "ChunkingError"
                }

            logger.info(f"Generated {len(chunks)} chunks")

            # 4. Generate embeddings
            logger.info("Generating embeddings...")
            chunk_texts = [chunk["raw_text"] for chunk in chunks]
            embeddings = self.embedding_handler.embed_batch(chunk_texts)

            # Add embeddings to chunks
            for chunk, embedding in zip(chunks, embeddings):
                chunk["embedding"] = embedding

            # 5. Store in vector database
            logger.info(f"Storing in collection '{collection_name}'")
            chunks_added = self.vector_store.add_documents(
                collection_name=collection_name,
                chunks=chunks,
                document_id=document_id,
                version=version
            )

            # Calculate statistics
            processing_time = time.time() - start_time
            file_size_mb = file_path.stat().st_size / (1024 * 1024)

            result = {
                "success": True,
                "document_id": document_id,
                "version": version,
                "filename": file_path.name,
                "chunks_created": chunks_added,
                "pages": extracted_metadata.get("page_count", 0),
                "file_size_mb": round(file_size_mb, 2),
                "ingested_at": datetime.utcnow().isoformat(),
                "processing_time_seconds": round(processing_time, 2),
                "collection_name": collection_name
            }

            logger.info(f"Ingestion complete: {chunks_added} chunks in {processing_time:.2f}s")

            return result

        except Exception as e:
            logger.error(f"Error ingesting document: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_type": type(e).__name__,
                "processing_time_seconds": round(time.time() - start_time, 2)
            }

    def _generate_document_id(self, file_path: Path) -> str:
        """Generate a unique document ID based on file path and name"""
        # Use hash of absolute path + filename for consistency
        unique_string = f"{file_path.absolute()}::{file_path.name}"
        return hashlib.md5(unique_string.encode()).hexdigest()[:16]

    def _extract_document(self, file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text and metadata from document.

        Args:
            file_path: Path to document

        Returns:
            Tuple of (text, metadata dict)
        """
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self._extract_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_docx(file_path)
        elif extension in ['.txt', '.md']:
            return self._extract_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text and metadata from PDF using PyMuPDF.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (text, metadata)
        """
        try:
            doc = fitz.open(file_path)

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")

            text = "\n\n".join(text_parts)

            # Extract metadata
            metadata = {
                "page_count": len(doc),
                "file_type": "pdf"
            }

            # Try to get PDF metadata
            pdf_metadata = doc.metadata
            if pdf_metadata:
                if pdf_metadata.get("title"):
                    metadata["title"] = pdf_metadata["title"]
                if pdf_metadata.get("author"):
                    metadata["author"] = pdf_metadata["author"]
                if pdf_metadata.get("subject"):
                    metadata["subject"] = pdf_metadata["subject"]

            doc.close()

            logger.info(f"Extracted {len(text)} characters from {len(doc)} pages")

            return text, metadata

        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise

    def _extract_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text and metadata from Word document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (text, metadata)
        """
        try:
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            text = "\n\n".join(paragraphs)

            # Metadata
            metadata = {
                "page_count": len(doc.sections),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "file_type": "docx"
            }

            # Try to get document properties
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
            except:
                pass

            logger.info(f"Extracted {len(text)} characters from Word document")

            return text, metadata

        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise

    def _extract_text(self, file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text from plain text file.

        Args:
            file_path: Path to text file

        Returns:
            Tuple of (text, metadata)
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = file_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                text = file_path.read_text(encoding='latin-1')

            # Metadata
            metadata = {
                "file_type": file_path.suffix[1:],  # Remove dot
                "line_count": len(text.split('\n'))
            }

            logger.info(f"Extracted {len(text)} characters from text file")

            return text, metadata

        except Exception as e:
            logger.error(f"Error extracting text file: {e}")
            raise


if __name__ == "__main__":
    # Test document processor
    from src.config import load_config
    from src.vector_store import VectorStore
    from src.embeddings import EmbeddingHandler
    from src.chunking import SemanticChunker

    # Initialize components
    config = load_config()
    vector_store = VectorStore(config)
    embedding_handler = EmbeddingHandler(config)
    chunker = SemanticChunker(config)

    # Create processor
    processor = DocumentProcessor(
        config,
        vector_store,
        embedding_handler,
        chunker
    )

    print("Document processor test initialized")
    print(f"Supported file types: {processor.SUPPORTED_EXTENSIONS}")
